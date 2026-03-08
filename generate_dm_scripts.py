"""
Generate 150 personalized Instagram DM scripts matching dm-scripts-30-final.docx style.
Reads new_outreach.csv for business data, writes a single Word doc.
"""

import csv
import re
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

CSV_PATH = "new_outreach.csv"
OUTPUT   = "dm-scripts-150.docx"


# ── Business type inference ────────────────────────────────────────────────────

def infer_type(name):
    n = name.lower()
    if any(x in n for x in ["infusion", "iv lounge", "iv hydration", "drip", "hydration"]):
        return "infusion"
    if any(x in n for x in ["laser hair", "laser skin", "laser & medical", "laser and medical"]):
        return "laser"
    if "dermatology" in n:
        return "dermatology"
    if "wellness" in n and any(x in n for x in ["medical", "med", "health"]):
        return "wellness"
    if "hyper wellness" in n or "restore" in n:
        return "wellness"
    if any(x in n for x in ["med spa", "medspa", "medispa", "medical spa"]):
        return "medspa"
    if any(x in n for x in ["aesthetic", "aesthetics"]):
        return "aesthetics"
    if "spa" in n and "salon" in n:
        return "salon_spa"
    if any(x in n for x in ["day spa", "spa"]):
        return "spa"
    if any(x in n for x in ["health", "wellness"]):
        return "wellness"
    if any(x in n for x in ["medical", "med "]):
        return "medical"
    return "medspa"


# ── Short name helper ─────────────────────────────────────────────────────────

def short_name(name):
    """Return a natural, conversational short version of the business name."""
    # Strip parenthetical suffixes like "(formerly ...)"
    name = re.sub(r'\s*\(.*?\)', '', name).strip()
    # Strip common location suffixes
    for suffix in [
        " - Boca Raton", " - Delray Beach", " - West Palm Beach",
        " - Palm Beach Gardens", " - Coral Springs", " - Deerfield Beach",
        " - Parkland", " Boca Raton", " Delray Beach", " West Palm Beach",
        " Palm Beach Gardens", " Coral Springs", " Deerfield Beach",
        " - PGA", " PGA",
    ]:
        if name.lower().endswith(suffix.lower()):
            name = name[: -len(suffix)].strip()
    # If name has colon (e.g. "Zarreii Medical and Aesthetics: Dr. Name"), keep first part
    if ":" in name:
        name = name.split(":")[0].strip()
    # Shorten very long names to first meaningful segment
    # e.g. "4Ever Young Med Spa & Wellness Center" -> "4Ever Young"
    overlong_map = {
        "4ever young med spa & wellness center": "4Ever Young",
        "4ever young med spa wellness center": "4Ever Young",
        "delray laser and medical spa": "Delray Laser and Medical Spa",
        "well labs medspa coconut creek": "Well Labs",
        "five elements medical spa wellness center in deerfield beach": "Five Elements",
        "doctor bg med excellence health well-being and longevity": "Doctor BG Med",
        "apex aesthetics and wellness jason kwiatkowski pa-c": "Apex Aesthetics",
        "alluramd boca raton formerly the box med spa": "AlluraMD",
        "liquivida wellness center coral springs": "Liquivida Coral Springs",
        "liquivida wellness center palm beach gardens pga": "Liquivida Palm Beach Gardens",
        "prime iv hydration wellness delray beach": "Prime IV",
        "green wellness centre med spa": "Green Wellness Centre",
        "magic touch medical esthetician by sandra maltese": "Magic Touch",
        "bareface aesthetics facial spa skin treatment": "BareFace Aesthetics",
        "zarreii medical and aesthetics": "Zarreii Medical and Aesthetics",
        "ashley swain permanent makeup medspa": "Ashley Swain",
        "pulsar health dr roop monica gill": "Pulsar Health",
        "sakara spa formerly known as belly love spa": "Sakara Spa",
        "romei med spa laser hair removal": "Romei Med Spa",
        "medclub testosterone bioidentical hormones": "MedClub",
    }
    key = re.sub(r'[^a-z0-9\s]', '', name.lower()).strip()
    key = re.sub(r'\s+', ' ', key)
    for k, v in overlong_map.items():
        if key == k:
            return v
    return name


# ── DM templates ─────────────────────────────────────────────────────────────
#
# Openers (vary by index % 5)
# Bodies (vary by business type + review tier)
# CTAs (vary by index % 4)
#
# All end with ": [url]"

OPENERS = [
    lambda sn, city: f"I came across your page and",
    lambda sn, city: f"I found your page and",
    lambda sn, city: f"I stumbled on your profile and",
    lambda sn, city: f"I was looking at {sn} and",
    lambda sn, city: f"I came across your business and",
]

CTAS = [
    "Would love to know what you think: [url]",
    "Would love to hear your thoughts: [url]",
    "Take a look when you get a chance: [url]",
    "Take a look when you get a moment: [url]",
]

def fmt_reviews(reviews):
    """Format review count naturally: 1,234 for 4-digit, plain for less."""
    return f"{reviews:,}" if reviews >= 1000 else str(reviews)


def review_phrase(reviews, rating):
    """A short complimentary phrase about their numbers."""
    rv = fmt_reviews(reviews)
    if rating == 5.0:
        if reviews >= 1000:
            return f"a perfect {rating:.1f} across {rv} reviews genuinely stopped me — that is seriously rare"
        elif reviews >= 400:
            return f"a perfect {rating:.1f} across {rv} reviews honestly stopped me — that is really hard to maintain"
        else:
            return f"a perfect {rating:.1f} with {rv} reviews is really impressive"
    elif rating >= 4.9:
        if reviews >= 500:
            return f"{rating:.1f} stars across {rv} reviews is a standout reputation"
        elif reviews >= 200:
            return f"{rv} reviews at {rating:.1f} stars is a really strong reputation"
        else:
            return f"{rating:.1f} stars with {rv} reviews is really solid"
    elif rating >= 4.8:
        if reviews >= 500:
            return f"{rating:.1f} stars with {rv} reviews is a beautiful reputation"
        elif reviews >= 200:
            return f"{rating:.1f} stars across {rv} reviews shows a really loyal client base"
        else:
            return f"a solid {rating:.1f} with {rv} reviews shows a consistent client experience"
    elif rating >= 4.7:
        if reviews >= 400:
            return f"{rating:.1f} stars across {rv} reviews shows a really loyal community"
        else:
            return f"{rating:.1f} stars with {rv} reviews shows a solid patient reputation"
    elif rating >= 4.6:
        return f"{rating:.1f} stars and {rv} reviews is a solid reputation in a competitive market"
    elif rating >= 4.5:
        return f"{rating:.1f} stars across {rv} reviews gives you a good trust baseline"
    else:
        return f"{rv} reviews on Google shows an established practice"


def city_qualifier(city):
    """Return 'in {city}' or empty if city is blank."""
    return f" in {city}" if city else ""


def body_sentence(sn, city, btype, reviews, rating, idx):
    """The 'I put together a little audit...' sentence."""
    cq = city_qualifier(city)

    templates = {
        "infusion": [
            f"I put together a little audit for {sn} showing how other infusion and wellness practices are increasing session frequency automatically.",
            f"I put together a little audit for {sn} showing how similar IV and wellness practices are getting more return visits from clients between drip sessions.",
            f"I put together a little audit for {sn} showing how other wellness practices{cq} are increasing recurring bookings from their existing client base.",
        ],
        "laser": [
            f"I put together a little audit for {sn} showing how other laser and aesthetic clinics are capturing more recurring visits after a series ends.",
            f"I put together a little audit for {sn} showing how similar laser practices are keeping clients coming back after a treatment plan finishes.",
            f"I put together a little audit for {sn} showing how other clinics{cq} are increasing rebooking after clients complete a laser package.",
        ],
        "dermatology": [
            f"I put together a little audit for {sn} showing how other cosmetic clinics are increasing return visits after treatment without any extra marketing effort.",
            f"I put together a little audit for {sn} showing how similar dermatology and cosmetic practices are getting more recurring appointments from patients they already have.",
            f"I put together a little audit for {sn} showing how other medical aesthetic practices{cq} are bringing patients back automatically after a procedure.",
        ],
        "wellness": [
            f"I put together a little audit for {sn} showing how other wellness practices are increasing return visits between treatment cycles.",
            f"I put together a little audit for {sn} showing how similar health and wellness practices are getting more recurring visits from clients they already have.",
            f"I put together a little audit for {sn} showing how other wellness centers{cq} are keeping clients engaged and coming back more consistently.",
        ],
        "medspa": [
            f"I put together a little audit for {sn} showing how other med spas are getting more return visits from the clients they already have.",
            f"I put together a little audit for {sn} showing how other local med spas are turning that kind of loyalty into even more recurring bookings.",
            f"I put together a little audit for {sn} showing how similar med spas{cq} are bringing clients back right when their treatment cycle is due.",
            f"I put together a little audit for {sn} showing how other med spas{cq} are increasing recurring revenue from the clients they already have.",
        ],
        "aesthetics": [
            f"I put together a little audit for {sn} showing how other aesthetic practices are unlocking more recurring bookings from the clients they already have.",
            f"I put together a little audit for {sn} showing how similar aesthetic studios are bringing clients back right when their treatment is due.",
            f"I put together a little audit for {sn} showing how other practices{cq} are increasing return visit frequency from their existing client base.",
        ],
        "salon_spa": [
            f"I put together a little audit for {sn} showing how other businesses with two service categories are getting more bookings by connecting both sides automatically.",
            f"I put together a little audit for {sn} showing how other salon and spa combos are getting more return visits from the clients they already have.",
        ],
        "spa": [
            f"I put together a little audit for {sn} showing how other spas are bringing past clients back more consistently.",
            f"I put together a little audit for {sn} showing how other local spas are increasing how often past clients come back without any extra work.",
            f"I put together a little audit for {sn} showing how similar day spas{cq} are increasing visit frequency from their existing client base.",
        ],
        "medical": [
            f"I put together a little audit for {sn} showing how other medical aesthetic practices are increasing patient return visits automatically.",
            f"I put together a little audit for {sn} showing how similar practices are increasing return visits from patients they already have.",
        ],
    }

    options = templates.get(btype, templates["medspa"])
    return options[idx % len(options)]


def make_dm(name, city, ig, reviews, rating, url, idx):
    sn     = short_name(name)
    btype  = infer_type(name)
    opener = OPENERS[idx % len(OPENERS)](sn, city)
    rp     = review_phrase(reviews, rating)
    cq     = city_qualifier(city)
    body   = body_sentence(sn, city, btype, reviews, rating, idx)
    cta    = CTAS[idx % len(CTAS)].replace("[url]", url)

    # Build the full DM
    return f"Hi! {opener} {rp}{cq}. {body} {cta}"


# ── Word doc builder ──────────────────────────────────────────────────────────

def build_doc(entries):
    doc = Document()

    # Title
    title = doc.add_heading("Instagram DM Scripts", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("150 Outreach Messages — Auralith Digital")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    doc.add_paragraph("")  # spacer

    for i, (num, name, city, ig, reviews, rating, url, dm_text) in enumerate(entries):
        # Numbered heading
        h = doc.add_paragraph()
        run = h.add_run(f"{num}. {name}")
        run.bold = True
        run.font.size = Pt(11)

        # Handle line
        handle_str = ig if ig else "Google / direct"
        hdl = doc.add_paragraph()
        r = hdl.add_run(f"Handle: {handle_str}")
        r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        r.font.size = Pt(10)

        # DM text
        dm_para = doc.add_paragraph(dm_text)
        dm_para.runs[0].font.size = Pt(10)

        # Spacer between entries (except last)
        if i < len(entries) - 1:
            doc.add_paragraph("")

    doc.save(OUTPUT)
    print(f"Saved {OUTPUT} ({len(entries)} DMs)")


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    os.chdir(os.path.dirname(os.path.abspath(__file__)))

    entries = []
    with open(CSV_PATH, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for i, row in enumerate(reader):
            name    = row["name"].strip()
            city    = row["city"].strip()
            ig      = row["ig"].strip()
            reviews = int(row["reviews"])
            url     = row["url"].strip()
            # derive rating from the HTML file if needed — but CSV doesn't have rating
            # Re-read from XLSX to get rating
            entries.append((i + 1, name, city, ig, reviews, url))

    # We need ratings — pull from XLSX
    import openpyxl

    def slugify(name):
        name = name.lower()
        name = re.sub(r'[^a-z0-9\s-]', '', name)
        name = re.sub(r'\s+', '-', name.strip())
        return re.sub(r'-+', '-', name)[:60]

    wb = openpyxl.load_workbook("MedSpaOutreach_v3.xlsx")
    lt = wb["Lead Tracker"]
    name_to_rating = {}
    for row in range(4, 800):
        n = lt.cell(row=row, column=1).value
        if not n: break
        try: rat = float(lt.cell(row=row, column=13).value or 0)
        except: rat = 0.0
        name_to_rating[str(n).strip().lower()] = rat

    final = []
    for num, name, city, ig, reviews, url in entries:
        rating = name_to_rating.get(name.lower(), 0.0)
        dm = make_dm(name, city, ig, reviews, rating, url, num - 1)
        final.append((num, name, city, ig, reviews, rating, url, dm))

    build_doc(final)

    # Preview first 3
    print("\nFirst 3 DMs:")
    for num, name, city, ig, reviews, rating, url, dm in final[:3]:
        print(f"\n{num}. {name}")
        print(f"   {dm}")


if __name__ == "__main__":
    main()
