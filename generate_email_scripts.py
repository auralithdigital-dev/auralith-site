"""
Generate 150 personalized cold email scripts for the med spa batch.
Reads new_outreach.csv + MedSpaOutreach_v3.xlsx for data.
Style: same voice as dm-scripts-150.docx but expanded to email format
(subject line + 3-4 short paragraphs + sign-off).
Output: email-scripts-150.docx
"""

import csv
import re
import os
import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

CSV_PATH = "new_outreach.csv"
XLSX_PATH = "MedSpaOutreach_v3.xlsx"
OUTPUT    = "email-scripts-150.docx"


# ── Reuse helpers from generate_dm_scripts ────────────────────────────────────

def slugify_name(name):
    name = name.lower()
    name = re.sub(r'[^a-z0-9\s-]', '', name)
    name = re.sub(r'\s+', '-', name.strip())
    return re.sub(r'-+', '-', name)[:60]

def short_name(name):
    name = re.sub(r'\s*\(.*?\)', '', name).strip()
    for suffix in [
        " - Boca Raton", " - Delray Beach", " - West Palm Beach",
        " - Palm Beach Gardens", " - Coral Springs", " - Deerfield Beach",
        " - Parkland", " Boca Raton", " Delray Beach", " West Palm Beach",
        " Palm Beach Gardens", " Coral Springs", " Deerfield Beach",
        " - PGA", " PGA",
    ]:
        if name.lower().endswith(suffix.lower()):
            name = name[: -len(suffix)].strip()
    if ":" in name:
        name = name.split(":")[0].strip()
    overlong = {
        "4ever young med spa & wellness center": "4Ever Young",
        "4ever young med spa wellness center":   "4Ever Young",
        "well labs medspa coconut creek":        "Well Labs",
        "five elements medical spa wellness center in deerfield beach": "Five Elements",
        "doctor bg med excellence health well-being and longevity": "Doctor BG Med",
        "apex aesthetics and wellness jason kwiatkowski pa-c": "Apex Aesthetics",
        "alluramd boca raton formerly the box med spa": "AlluraMD",
        "liquivida wellness center coral springs": "Liquivida Coral Springs",
        "liquivida wellness center palm beach gardens pga": "Liquivida Palm Beach Gardens",
        "prime iv hydration wellness delray beach": "Prime IV",
        "green wellness centre med spa": "Green Wellness Centre",
        "magic touch medical esthetician by sandra maltese": "Magic Touch",
        "bareface aesthetics facial spa skin treatment": "BareFace Aesthetics",
        "ashley swain permanent makeup medspa": "Ashley Swain",
        "pulsar health dr roop monica gill": "Pulsar Health",
        "sakara spa formerly known as belly love spa": "Sakara Spa",
        "romei med spa laser hair removal": "Romei Med Spa",
        "medclub testosterone bioidentical hormones": "MedClub",
        "delray laser and medical spa": "Delray Laser and Medical Spa",
        "zarreii medical and aesthetics": "Zarreii Medical and Aesthetics",
        "centro ddl bioesttica integral": "Centro DDL",
        "centro ddl bioestetica integral": "Centro DDL",
        "invictus medical and wellness group": "Invictus Medical",
        "athena med spa and salon": "Athena",
        "cr8 health wellness aesthetics": "CR8 Health",
        "five elements medical spa wellness center in deerfield beach": "Five Elements",
        "well labs medspa coconut creek formerly hello gorgeous": "Well Labs",
    }
    key = re.sub(r'[^a-z0-9\s]', '', name.lower()).strip()
    key = re.sub(r'\s+', ' ', key)
    return overlong.get(key, name)

def clean_name(name):
    """Strip legal suffixes."""
    return re.sub(r'\s*,?\s*\b(inc\.?|llc\.?|corp\.?|ltd\.?)\s*$', '', name, flags=re.IGNORECASE).strip()

def fmt_reviews(reviews):
    return f"{reviews:,}" if reviews >= 1000 else str(reviews)

def infer_type(name):
    n = name.lower()
    if any(x in n for x in ["infusion", "iv lounge", "iv hydration", "drip", "hydration"]):
        return "infusion"
    if any(x in n for x in ["laser hair", "laser skin", "laser & medical", "laser and medical"]):
        return "laser"
    if "dermatology" in n:
        return "dermatology"
    if "spa" in n and "salon" in n:
        return "salon_spa"
    if "wellness" in n and any(x in n for x in ["medical", "med", "health"]):
        return "wellness"
    if "hyper wellness" in n or "restore" in n:
        return "wellness"
    if any(x in n for x in ["med spa", "medspa", "medispa", "medical spa"]):
        return "medspa"
    if any(x in n for x in ["aesthetic", "aesthetics"]):
        return "aesthetics"
    if any(x in n for x in ["day spa", "spa"]):
        return "spa"
    if any(x in n for x in ["health", "wellness"]):
        return "wellness"
    return "medspa"


# ── Subject lines (5 variants, rotate by idx % 5) ────────────────────────────

def subject_line(name, city, reviews, rating, idx):
    sn = short_name(name)
    cn = clean_name(sn)
    variants = [
        f"{cn} — a thought on client retention",
        f"quick note for {cn}",
        f"noticed something about {cn}",
        f"{cn} / your {city} clients" if city else f"{cn} / your clients",
        f"idea for {cn}",
    ]
    return variants[idx % len(variants)]


# ── Opening lines (5 variants) ────────────────────────────────────────────────

def opening(name, city, reviews, rating, idx):
    sn  = short_name(name)
    rv  = fmt_reviews(reviews)
    cq  = f" in {city}" if city else ""

    if rating == 5.0 and reviews >= 400:
        openers = [
            f"I came across {sn} while researching med spas{cq} and a perfect {rating:.1f}-star rating across {rv} reviews honestly stopped me.",
            f"I was looking at practices{cq} and {sn}'s perfect {rating:.1f} across {rv} reviews is genuinely rare — that kind of reputation takes years to earn.",
            f"I found {sn} online and a perfect {rating:.1f} with {rv} reviews{cq} is really something. That level of client loyalty stands out.",
            f"I came across your practice and {rv} reviews at a perfect {rating:.1f} stars{cq} is the kind of reputation most med spas spend years trying to build.",
            f"I was looking at med spas{cq} and {sn}'s {rv} reviews at {rating:.1f} stars genuinely stopped me — that is not easy to maintain.",
        ]
    elif rating >= 4.8 and reviews >= 300:
        openers = [
            f"I came across {sn} and {rv} reviews at {rating:.1f} stars{cq} is a really strong reputation.",
            f"I found your practice online — {rating:.1f} stars across {rv} reviews{cq} shows a client base that clearly trusts you.",
            f"I was looking at practices{cq} and {sn}'s {rating:.1f}-star rating with {rv} reviews stood out to me.",
            f"I came across your page and {rv} reviews at {rating:.1f} stars{cq} is a solid foundation to build on.",
            f"I found {sn} while researching local practices — {rating:.1f} stars and {rv} reviews{cq} is a really consistent track record.",
        ]
    elif reviews >= 200:
        openers = [
            f"I came across {sn} and {rv} reviews at {rating:.1f} stars{cq} shows an established, trusted practice.",
            f"I found your practice online — {rv} reviews at {rating:.1f} stars{cq} gives you a solid trust baseline.",
            f"I was looking at practices{cq} and {sn}'s {rv} Google reviews caught my attention.",
            f"I came across your page and {rating:.1f} stars across {rv} reviews{cq} is a reputation worth protecting and growing.",
            f"I found {sn} online — {rv} reviews at {rating:.1f} stars{cq} shows a client base you've worked hard to build.",
        ]
    else:
        openers = [
            f"I came across {sn} online — {rv} reviews at {rating:.1f} stars{cq} shows a growing practice.",
            f"I found your practice and {rv} Google reviews at {rating:.1f} stars{cq} is a solid start.",
            f"I was looking at practices{cq} and {sn}'s {rating:.1f}-star rating across {rv} reviews stood out.",
            f"I came across your page — {rv} reviews at {rating:.1f} stars{cq} is a good reputation to build on.",
            f"I found {sn} while researching local practices — {rv} reviews at {rating:.1f} stars{cq} shows consistency.",
        ]

    return openers[idx % len(openers)]


# ── Observation line (the gap / opportunity) ──────────────────────────────────

def observation(name, city, reviews, btype, idx):
    sn = short_name(name)
    cq = f" in {city}" if city else ""

    obs_map = {
        "infusion": [
            f"Practices like {sn} with that kind of reputation usually have a gap in session re-engagement — clients who finish a drip series and don't hear back until they decide to rebook on their own.",
            f"The gap I see most often with infusion practices is the window between sessions — clients who had a great experience but don't get a nudge to come back before they've already moved on.",
            f"Most infusion and IV practices{cq} lose a portion of clients simply because there's no automated touchpoint after a session ends.",
        ],
        "laser": [
            f"Most laser practices{cq} have a window after a series ends where clients drift — no automated follow-up means they forget to book their next round until it's months later.",
            f"The gap I see most often with laser clinics is after a series finishes — clients don't always self-schedule the next phase without a nudge.",
            f"Practices like {sn} typically have a drop-off between laser series — clients complete a round and don't hear from the clinic until they decide to reach out.",
        ],
        "dermatology": [
            f"Even with a strong reputation, most cosmetic clinics{cq} have a portion of past patients who haven't returned since their last treatment and aren't getting any follow-up to bring them back.",
            f"The gap I see most often with cosmetic practices is after a procedure — clients leave happy but don't get a proactive touchpoint to schedule a follow-up.",
            f"Most cosmetic clinics{cq} rely on patients to self-schedule, which means a portion of happy clients quietly drift to other options without realizing it.",
        ],
        "wellness": [
            f"Most wellness practices{cq} have a gap between treatment cycles — clients who finish a protocol and don't hear from the practice until they've already looked elsewhere.",
            f"The gap I see most often with wellness centers is the period between visits — clients who had great results but don't get a proactive touchpoint to bring them back.",
            f"Even with a strong reputation, practices like {sn} typically have a portion of clients who go inactive between cycles without any automated re-engagement.",
        ],
        "medspa": [
            f"Most med spas{cq} have a gap between treatments — clients who finish a Botox or filler cycle and don't hear from the practice until they decide to rebook on their own.",
            f"The gap I see most often with practices like {sn} is in client re-engagement — past clients who had a great experience but don't get a proactive nudge to come back.",
            f"Even with a reputation like {sn}'s, most practices{cq} have a portion of clients who go quietly inactive between treatment cycles without any automated follow-up.",
        ],
        "aesthetics": [
            f"Most aesthetic practices{cq} have a window after a treatment where clients could easily rebook — but without an automated nudge, a portion of them don't come back until they've already checked out a competitor.",
            f"The gap I see most often with aesthetic studios like {sn} is in recurring visits — clients love the result but don't always self-schedule their next session.",
            f"Even with strong reviews, most aesthetic practices{cq} have clients who go inactive between appointments without any proactive outreach to bring them back.",
        ],
        "salon_spa": [
            f"Businesses that offer both salon and spa services{cq} often have clients who use one side and don't realize what the other side offers — a missed cross-booking opportunity every appointment.",
            f"The gap I see most often with salon and spa combos is in cross-service booking — clients who book hair and never hear about the spa side, or vice versa.",
        ],
        "spa": [
            f"Most day spas{cq} have a gap in returning client outreach — past guests who had a great experience and don't hear back until they decide to look for a spa again.",
            f"The gap I see most often with spas like {sn} is in re-engagement — clients who visited once or twice and quietly drift without any automated follow-up to bring them back.",
            f"Even with strong reviews, most spas{cq} have a portion of past guests who haven't returned recently and aren't getting any proactive outreach.",
        ],
        "medical": [
            f"Most medical aesthetic practices{cq} have a gap in patient re-engagement after a procedure — patients who had great results and don't get a proactive follow-up before they've already started looking elsewhere.",
            f"The gap I see most often is after a treatment ends — patients leave happy but don't hear from the practice again until they decide to rebook on their own.",
        ],
    }

    options = obs_map.get(btype, obs_map["medspa"])
    return options[idx % len(options)]


# ── Audit pitch line ──────────────────────────────────────────────────────────

def audit_pitch(name, city, reviews, btype, url, idx):
    sn = short_name(name)
    cq = f" in {city}" if city else ""

    pitches = [
        f"I put together a short audit for {sn} showing specifically how other practices{cq} are handling this automatically. You can see what I found here: {url}",
        f"I put together a quick breakdown for {sn} showing how similar practices{cq} are solving this without adding any manual work. Take a look when you get a minute: {url}",
        f"I put together a short audit showing the specific gaps I noticed for {sn} and how other {city + ' ' if city else ''}practices are closing them automatically: {url}",
        f"I put together a quick audit for {sn} — it shows what's likely happening with your inactive clients and how similar practices{cq} are fixing it: {url}",
        f"I put together a short breakdown for {sn} showing how other practices with a reputation like yours are turning past clients into recurring visits automatically: {url}",
    ]
    return pitches[idx % len(pitches)]


# ── Sign-off lines (vary) ─────────────────────────────────────────────────────

SIGNOFFS = [
    "Happy to answer any questions.\n\nBest,\nPyetra",
    "No pitch — just the findings.\n\nBest,\nPyetra",
    "Totally fine if it's not relevant, just didn't want it to go unseen.\n\nBest,\nPyetra",
    "Let me know if any of it resonates.\n\nBest,\nPyetra",
    "No pressure at all — just thought it was worth sharing.\n\nBest,\nPyetra",
]


# ── Assemble full email ───────────────────────────────────────────────────────

def make_email(name, city, ig, reviews, rating, url, idx):
    btype   = infer_type(name)
    subject = subject_line(name, city, reviews, rating, idx)
    op      = opening(name, city, reviews, rating, idx)
    obs     = observation(name, city, reviews, btype, idx)
    pitch   = audit_pitch(name, city, reviews, btype, url, idx)
    signoff = SIGNOFFS[idx % len(SIGNOFFS)]

    body = f"{op}\n\n{obs}\n\n{pitch}\n\n{signoff}"
    return subject, body


# ── Word doc builder ──────────────────────────────────────────────────────────

def build_doc(entries):
    doc = Document()

    # Title
    title = doc.add_heading("Email Scripts", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("150 Cold Email Scripts — Auralith Digital")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    doc.add_paragraph("")

    for i, (num, name, city, ig, reviews, rating, url, subject, body) in enumerate(entries):
        # Numbered heading
        h = doc.add_paragraph()
        r = h.add_run(f"{num}. {name}")
        r.bold = True
        r.font.size = Pt(11)

        # Handle / email target
        handle_str = ig if ig else "Google / direct"
        hdl = doc.add_paragraph()
        hr = hdl.add_run(f"Handle: {handle_str}")
        hr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        hr.font.size = Pt(10)

        # Subject line
        subj_para = doc.add_paragraph()
        subj_para.add_run("Subject: ").bold = True
        subj_run = subj_para.add_run(subject)
        subj_run.font.size = Pt(10)

        # Body
        for line in body.split("\n"):
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].font.size = Pt(10)

        if i < len(entries) - 1:
            doc.add_paragraph("")

    doc.save(OUTPUT)
    print(f"Saved {OUTPUT} ({len(entries)} emails)")


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    os.chdir(os.path.dirname(os.path.abspath(__file__)))

    # Load ratings from XLSX
    wb = openpyxl.load_workbook(XLSX_PATH)
    lt = wb["Lead Tracker"]
    name_to_rating = {}
    for row in range(4, 800):
        n = lt.cell(row=row, column=1).value
        if not n:
            break
        try:
            rat = float(lt.cell(row=row, column=13).value or 0)
        except:
            rat = 0.0
        name_to_rating[str(n).strip().lower()] = rat

    # Load CSV
    entries = []
    with open(CSV_PATH, newline="", encoding="utf-8") as f:
        for i, row in enumerate(csv.DictReader(f)):
            name    = row["name"].strip()
            city    = row["city"].strip()
            ig      = row["ig"].strip()
            reviews = int(row["reviews"])
            url     = row["url"].strip()
            rating  = name_to_rating.get(name.lower(), 0.0)
            subject, body = make_email(name, city, ig, reviews, rating, url, i)
            entries.append((i + 1, name, city, ig, reviews, rating, url, subject, body))

    build_doc(entries)

    print("\nSample (entries 1, 15, 50, 100, 150):")
    for idx in [0, 14, 49, 99, 149]:
        num, name, city, ig, reviews, rating, url, subject, body = entries[idx]
        print(f"\n{'─'*60}")
        print(f"{num}. {name}  [{reviews} reviews, {rating}★]")
        print(f"Subject: {subject}")
        print(body)


if __name__ == "__main__":
    main()
